"""
generate.py

Generates DOCX documents from extracted form data.

Supported document types:
- Office Leave Application
- University Leave Application
- Complaint Letter
"""

from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Pt

# Import schemas
from src.extraction.schemas import (
    LeaveApplicationOfficeForm,
    LeaveApplicationUniversityForm,
    ComplaintLetterForm,
)
from src.utils.date_utils import format_date



def _add_line(
    doc: Document,
    text,
    bold: bool = False,
    size: int = 11,
):
    """
    Adds a single paragraph to the document.

    If text is None or empty, nothing is added.
    """

    if not text:
        return

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)


def write_body_paragraph_office(
    form: LeaveApplicationOfficeForm,
) -> str:
    """
    Generates the body paragraph for an office leave application.
    """

    start_display = format_date(form.start_date)

    return (
        f"I am writing to respectfully request "
        f"{form.duration_days} day(s) of "
        f"{form.leave_type} leave starting "
        f"{start_display}, on account of "
        f"{form.reason}. "
        f"I will ensure that all pending tasks are "
        f"properly handed over before my leave begins."
    )


def _generate_office(
    form: LeaveApplicationOfficeForm,
    output_path: Path,
) -> Path:
    """
    Generate an Office Leave Application DOCX.
    """

    doc = Document()

    today_display = format_date(date.today().isoformat())

    duration_display = (
        f"{form.duration_days} day"
        f"{'s' if form.duration_days != 1 else ''}"
    )


    _add_line(doc, form.applicant_name, bold=True)

    department_line = (
        f"{form.applicant_designation}, {form.department}"
        if form.department
        else form.applicant_designation
    )

    _add_line(doc, department_line)

    if form.employee_id:
        _add_line(doc, f"Employee ID: {form.employee_id}")

    if form.contact_number:
        _add_line(doc, f"Contact: {form.contact_number}")

    _add_line(doc, f"Date: {today_display}")


    doc.add_paragraph()

    _add_line(doc, f"To: {form.recipient_name}")
    _add_line(doc, form.recipient_designation)
    _add_line(doc, form.company_name)


    doc.add_paragraph()

    leave_label = (
        form.leave_type.capitalize()
        if form.leave_type
        else "Leave"
    )

    _add_line(
        doc,
        f"Subject: Application for {leave_label} Leave ({duration_display})",
        bold=True,
    )


    doc.add_paragraph()

    _add_line(doc, "Respected Sir/Madam,")


    doc.add_paragraph()

    _add_line(
        doc,
        write_body_paragraph_office(form),
    )

    doc.add_paragraph()

    _add_line(
        doc,
        "I shall be grateful for your kind approval. "
        "Thank you for your consideration."
    )
    doc.add_paragraph()

    _add_line(doc, "Yours sincerely,")

    _add_line(doc, form.applicant_name)

    _add_line(doc, form.applicant_designation)

    # Save file

    doc.save(output_path)

    return output_path

def write_body_paragraph_university(
    form: LeaveApplicationUniversityForm,
) -> str:
    """
    Generates the body paragraph for a university leave application.
    """

    start_display = format_date(form.start_date)

    return (
        f"I respectfully submit that due to {form.reason}, "
        f"I am unable to attend my classes starting "
        f"{start_display}. Therefore, I kindly request "
        f"leave for {form.duration_days} day(s)."
    )


def _generate_university(
    form: LeaveApplicationUniversityForm,
    output_path: Path,
) -> Path:
    """
    Generate a University Leave Application DOCX.
    """

    doc = Document()

    today_display = format_date(date.today().isoformat())

    duration_display = (
        f"{form.duration_days} day"
        f"{'s' if form.duration_days != 1 else ''}"
    )


    _add_line(doc, f"Date: {today_display}")


    doc.add_paragraph()

    _add_line(doc, f"To: {form.recipient_designation}")

    if form.department:
        _add_line(
            doc,
            f"Department of {form.department}",
        )

    _add_line(doc, form.institution_name)

    doc.add_paragraph()

    _add_line(
        doc,
        f"Subject: Application for Leave ({duration_display})",
        bold=True,
    )


    doc.add_paragraph()

    salutation = form.recipient_salutation or "Sir/Madam"

    _add_line(
        doc,
        f"Respected {salutation},",
    )

    doc.add_paragraph()

    _add_line(
        doc,
        write_body_paragraph_university(form),
    )

    doc.add_paragraph()

    _add_line(
        doc,
        "Kindly grant me leave for the above-mentioned period. "
        "I assure you that I will cover all missed lectures, "
        "assignments, and coursework after returning."
    )
    doc.add_paragraph()

    _add_line(doc, "Yours obediently,")

    _add_line(doc, form.student_name)

    # Program + Semester

    if form.semester:
        _add_line(
            doc,
            f"{form.program}, Semester {form.semester}",
        )
    else:
        _add_line(
            doc,
            form.program,
        )

    _add_line(
        doc,
        f"Roll No: {form.roll_number}",
    )

    if form.institution_name:
        _add_line(
            doc,
            form.institution_name,
        )
    doc.save(output_path)

    return output_path

def write_body_paragraph_complaint(
    form: ComplaintLetterForm,
) -> str:
    """
    Generates the main body paragraph for a complaint letter.
    """

    return (
        f"I wish to bring to your attention that "
        f"{form.issue_description}. "
        f"This issue has caused significant inconvenience, "
        f"and I kindly request that appropriate action be taken "
        f"to resolve the matter as soon as possible."
    )


def _generate_complaint(
    form: ComplaintLetterForm,
    output_path: Path,
) -> Path:
    """
    Generate a Complaint Letter DOCX.
    """

    doc = Document()

    today_display = format_date(date.today().isoformat())

    _add_line(doc, form.complainant_name, bold=True)

    if form.address:
        _add_line(doc, form.address)

    if form.contact_number:
        _add_line(doc, f"Contact: {form.contact_number}")

    if form.reference_number:
        _add_line(
            doc,
            f"Reference No: {form.reference_number}",
        )

    _add_line(doc, f"Date: {today_display}")


    doc.add_paragraph()

    _add_line(
        doc,
        f"To: {form.recipient_designation}",
    )

    _add_line(doc, form.organization_name)

    if form.organization_address:
        _add_line(doc, form.organization_address)


    doc.add_paragraph()

    _add_line(
        doc,
        f"Subject: Complaint Regarding {form.complaint_subject}",
        bold=True,
    )


    doc.add_paragraph()

    _add_line(doc, "Dear Sir/Madam,")

    doc.add_paragraph()

    _add_line(
        doc,
        write_body_paragraph_complaint(form),
    )

    doc.add_paragraph()

    _add_line(
        doc,
        "I would appreciate your prompt attention "
        "to this matter and look forward to "
        "a suitable resolution at the earliest."
    )

    doc.add_paragraph()

    _add_line(doc, "Yours faithfully,")

    _add_line(doc, form.complainant_name)

    # Save document

    doc.save(output_path)

    return output_path


_GENERATORS = {
    "leave_application_office": _generate_office,
    "leave_application_university": _generate_university,
    "complaint_letter": _generate_complaint,
}


def generate(
    form,
    output_path: str | Path,
) -> Path:
    """
    Main entry point.

    Parameters
    ----------
    form
        Pydantic form object.

    output_path
        Where the DOCX should be saved.

    Returns
    -------
    Path
        Path to the generated document.
    """

    if not form.is_complete():
        raise ValueError(
            f"Cannot generate document. "
            f"Missing required fields: {form.missing_fields}"
        )

    generator = _GENERATORS.get(form.document_type)

    if generator is None:
        raise ValueError(
            f"No generator implemented for "
            f"{form.document_type}"
        )

    return generator(
        form,
        Path(output_path),
    )


if __name__ == "__main__":

    test_form = LeaveApplicationOfficeForm(
        applicant_name="Ayesha Niazi",
        applicant_designation="Software Engineer",
        recipient_name="Ahmed Khan",
        recipient_designation="HR Manager",
        company_name="TechSol Pvt Ltd",
        leave_type="casual",
        start_date="2026-08-26",
        duration_days=3,
        reason="family event",
    )

    output = generate(
        test_form,
        "sample_leave_application.docx",
    )

    print(f"Generated: {output}")